#!python


"""
THIS IS NOT ONE OF MY STRATEGIES NOR IS PROFITABLE - IT'S JUST TO SHOW YOU HOW I BUILT THE BACKTESTING
"""


class Engine:

    def __init__(self):

        import os
        import oandapyV20
        from oandapyV20 import API

        password = os.environ.get("Your_API_Key")
        user_name = os.environ.get("Your_Oanda_Account")

        # IF YOU USE A DEMO ACCOUNT REMOVE , environment="live"
        self.client = oandapyV20.API(access_token=password, environment="live")

        self.open_prices = []
        self.high_prices = []
        self.low_prices = []
        self.close_prices = []
        self.date_time = []

        self.forex_pairs = []
        self.forex_pairs_no_hkd = []

    # ****************************************************************************************************************

    def forex_pairs_names(self):

        self.forex_pairs_no_hkd = ["AUD_CAD", "AUD_CHF", "AUD_JPY", "AUD_NZD", "AUD_SGD", "AUD_USD",
                                   "CAD_CHF", "CAD_JPY", "CAD_SGD", "CHF_JPY",
                                   "EUR_AUD", "EUR_CAD", "EUR_CHF", "EUR_JPY", "EUR_GBP", "EUR_NZD", "EUR_SGD",
                                   "EUR_USD",
                                   "GBP_AUD", "GBP_CAD", "GBP_CHF", "GBP_JPY", "GBP_NZD", "GBP_SGD", "GBP_USD",
                                   "NZD_CAD", "NZD_CHF", "NZD_JPY", "NZD_SGD", "NZD_USD", "SGD_CHF", "SGD_JPY",
                                   "USD_CAD", "USD_CHF", "USD_JPY", "USD_SGD"]

        self.forex_pairs = ["AUD_HKD", "CAD_HKD", "CHF_HKD", "EUR_HKD", "GBP_HKD",
                            "NZD_HKD", "SGD_HKD"]

        self.forex_pairs.extend(self.forex_pairs_no_hkd)

    # *****************************************************************************************************************

    def historical_klines(self, forex_pair, num_of_candles, interval):

        import oandapyV20.endpoints.instruments as instruments

        parameters = {"count": num_of_candles, "granularity": interval}
        r = instruments.InstrumentsCandles(instrument=forex_pair, params=parameters)
        self.client.request(r)
        forex_pair_data_raw = r.response

        # OHLC
        flag_historical = True
        idx_historical = 0

        while flag_historical:

            self.open_prices.append(float(forex_pair_data_raw["candles"][idx_historical]["mid"]["o"]))
            self.high_prices.append(float(forex_pair_data_raw["candles"][idx_historical]["mid"]["h"]))
            self.low_prices.append(float(forex_pair_data_raw["candles"][idx_historical]["mid"]["l"]))
            self.close_prices.append(float(forex_pair_data_raw["candles"][idx_historical]["mid"]["c"]))

            self.date_time.append(forex_pair_data_raw["candles"][idx_historical]["time"][: -11])

            if idx_historical == num_of_candles - 1:
                flag_historical = False

            idx_historical += 1


# ! *******************************************************************************************************************


class Strategy(Engine):
    from_zero_trades_all_pairs = 0
    from_zero_trades_w_profit_all_pairs = 0
    from_zero_trades_w_loss_all_pairs = 0
    from_zero_profit_percentage_all_pairs = 0

    from_hundred_trades_all_pairs = 0
    from_hundred_trades_w_profit_all_pairs = 0
    from_hundred_trades_w_loss_all_pairs = 0
    from_hundred_profit_percentage_all_pairs = 0

    results = {}

    in_depth_res_long = []
    in_depth_res_short = []

    results_w_combos = []

    def __init__(self):

        Engine.__init__(self)

        # SINGLE PAIR
        self.ema_60 = []
        self.ema_20 = []

        self.close_prices_60 = []
        self.open_prices_60 = []
        self.low_prices_60 = []
        self.high_prices_60 = []
        self.ema_20_60 = []
        self.date_time_60 = []

        self.from_zero_data = []
        self.from_zero_trades = 0
        self.from_zero_trades_w_profit = 0
        self.from_zero_trades_w_loss = 0
        self.from_zero_profit_minus_loss = 0

        self.from_hundred_data = []
        self.from_hundred_trades = 0
        self.from_hundred_trades_w_profit = 0
        self.from_hundred_trades_w_loss = 0
        self.from_hundred_profit_minus_loss = 0

    # ***************************************************************************************************************

    def backtesting_single_pair(self, forex_pair, num_of_candles, interval,
                                take_profit, stop_loss, price_gap):

        self.__init__()
        self.historical_klines(forex_pair, num_of_candles, interval)

        assert len(self.close_prices) > 60, "NOT ENOUGH CANDLES"

        # EMA 60
        previous_ema = round(sum(self.close_prices[: 60]) / 60, 5)
        k_ema_60 = 2 / (60 + 1)
        flag_ema_60 = True
        idx_ema_60 = 60

        while flag_ema_60:

            current_ema = round((self.close_prices[idx_ema_60] * k_ema_60) + (previous_ema * (1 - k_ema_60)), 5)
            self.ema_60.append(current_ema)
            previous_ema = current_ema

            idx_ema_60 += 1

            if idx_ema_60 > len(self.close_prices) - 1:
                flag_ema_60 = False

        # EMA 20
        previous_ema = round(sum(self.close_prices[: 20]) / 20, 5)
        k_ema_20 = 2 / (20 + 1)
        flag_ema_20 = True
        idx_ema_20 = 20

        while flag_ema_20:

            current_ema = round((self.close_prices[idx_ema_20] * k_ema_20) + (previous_ema * (1 - k_ema_20)), 5)
            self.ema_20.append(current_ema)
            previous_ema = current_ema

            idx_ema_20 += 1

            if idx_ema_20 > len(self.close_prices) - 1:
                flag_ema_20 = False

        # ******************************************** STRATEGY ******************************************************
        len_diff = len(self.ema_60)
        self.close_prices_60 = self.close_prices[-len_diff:]
        self.open_prices_60 = self.open_prices[-len_diff:]
        self.low_prices_60 = self.low_prices[-len_diff:]
        self.high_prices_60 = self.high_prices[-len_diff:]
        self.ema_20_60 = self.ema_20[-len_diff:]
        self.date_time_60 = self.date_time[-len_diff:]

        if "JPY" in forex_pair:
            take_profit_pair = take_profit * 0.01
            stop_loss_pair = stop_loss * 0.01
            price_gap = price_gap * 0.01
            spread = 3.5 * 0.01
        elif "HKD" in forex_pair:
            take_profit_pair = take_profit * 0.001
            stop_loss_pair = stop_loss * 0.001
            price_gap = price_gap * 0.001
            spread = 3.5 * 0.001
        else:
            take_profit_pair = take_profit * 0.0001
            stop_loss_pair = stop_loss * 0.0001
            price_gap = price_gap * 0.0001
            spread = 3.5 * 0.0001

        # ******************************************  LONG ******************************************************
        flag = True
        idx = 5

        while flag:

            if idx > len(self.ema_60) - 1:

                flag = False

            elif self.ema_20_60[idx] <= self.close_prices_60[idx] >= self.ema_60[idx] and \
                    self.close_prices_60[idx] - self.open_prices_60[idx - 5] >= -price_gap:

                open_trade = self.close_prices_60[idx] - spread

                self.from_zero_data.append(("O", self.date_time_60[idx], idx, self.close_prices_60[idx],
                                            self.ema_20_60[idx], self.open_prices_60[idx - 5],
                                            round(self.close_prices_60[idx] - self.open_prices_60[idx - 5], 5)))
                self.from_zero_trades += 1

                inner_flag = True

                while inner_flag:

                    idx += 1

                    if idx > len(self.ema_60) - 1:
                        flag = False
                        inner_flag = False

                    elif self.low_prices_60[idx] - open_trade <= stop_loss_pair:

                        self.from_zero_data.append(
                            ("L", self.date_time_60[idx], idx, self.close_prices_60[idx], self.ema_20_60[idx]))

                        self.from_zero_trades_w_loss += 1

                        inner_flag = False

                    elif self.high_prices_60[idx] - open_trade >= take_profit_pair:

                        self.from_zero_data.append(
                            ("P", self.date_time_60[idx], idx, self.close_prices_60[idx], self.ema_20_60[idx]))

                        self.from_zero_trades_w_profit += 1

                        inner_flag = False

                    else:
                        continue

            idx += 1

        self.from_zero_profit_minus_loss = self.from_zero_trades_w_profit - self.from_zero_trades_w_loss

        # ************************************** SHORT *********************************************************
        flag = True
        idx = 5

        while flag:

            if idx > len(self.ema_60) - 1:

                flag = False

            elif self.ema_20_60[idx] >= self.close_prices_60[idx] <= self.ema_60[idx] and \
                    self.close_prices_60[idx] - self.open_prices_60[idx - 5] <= price_gap:

                open_trade = self.close_prices_60[idx] - spread

                self.from_hundred_data.append(("O", self.date_time_60[idx], idx, self.close_prices_60[idx],
                                               self.ema_20_60[idx], self.open_prices_60[idx - 5],
                                               round(self.close_prices_60[idx] - self.open_prices_60[idx - 5], 5)))
                self.from_hundred_trades += 1

                inner_flag = True

                while inner_flag:

                    idx += 1

                    if idx > len(self.ema_60) - 1:

                        flag = False
                        inner_flag = False

                    elif self.high_prices_60[idx] - open_trade >= -stop_loss_pair:

                        self.from_hundred_data.append(
                            ("L", self.date_time_60[idx], idx, self.close_prices_60[idx], self.ema_20_60[idx]))

                        self.from_hundred_trades_w_loss += 1

                        inner_flag = False

                    elif self.low_prices_60[idx] - open_trade <= -take_profit_pair:

                        self.from_hundred_data.append(
                            ("P", self.date_time_60[idx], idx, self.close_prices_60[idx], self.ema_20_60[idx]))

                        self.from_hundred_trades_w_profit += 1

                        inner_flag = False

                    else:
                        continue
            idx += 1

        self.from_hundred_profit_minus_loss = self.from_hundred_trades_w_profit - self.from_hundred_trades_w_loss

    # ************************************************* ALL PAIRS *************************************************

    def backtesting_all_pair(self, num_of_candles, interval, take_profit, stop_loss, price_gap):

        import time

        self.forex_pairs_names()

        for p in self.forex_pairs:

            try:

                self.backtesting_single_pair(p, num_of_candles, interval,
                                             take_profit, stop_loss, price_gap)

                self.from_zero_trades_all_pairs += self.from_zero_trades
                self.from_zero_trades_w_profit_all_pairs += self.from_zero_trades_w_profit
                self.from_zero_trades_w_loss_all_pairs += self.from_zero_trades_w_loss
                self.in_depth_res_long.append((p,))

                self.from_hundred_trades_all_pairs += self.from_hundred_trades
                self.from_hundred_trades_w_profit_all_pairs += self.from_hundred_trades_w_profit
                self.from_hundred_trades_w_loss_all_pairs += self.from_hundred_trades_w_loss
                self.in_depth_res_short.append((p, self.from_hundred_trades, self.from_hundred_profit_minus_loss))

                time.sleep(1)

            except:

                print("ERROR WITH", p)

                continue

        win_percentage_short = round(((100 / (self.from_hundred_trades_w_profit_all_pairs +
                                              self.from_hundred_trades_w_loss_all_pairs)) *
                                      self.from_hundred_trades_w_profit_all_pairs), 2)

        win_percentage_long = round(((100 / (self.from_zero_trades_w_profit_all_pairs +
                                             self.from_zero_trades_w_loss_all_pairs)) *
                                     self.from_zero_trades_w_profit_all_pairs), 2)

        winning_trades_short = self.from_hundred_trades_w_profit_all_pairs - self.from_hundred_trades_w_loss_all_pairs

        winning_trades_long = self.from_zero_trades_w_profit_all_pairs - self.from_zero_trades_w_loss_all_pairs

        self.results = {"TP / SL": [take_profit, stop_loss],
                        "TOT LONG TRADES": self.from_zero_trades_all_pairs,
                        "WIN - LOSS LONG": winning_trades_long,
                        "WINNING PERCENTAGE LONG": win_percentage_long,
                        "TOT SHORT TRADES": self.from_hundred_trades_all_pairs,
                        "WIN - LOSS SHORT": winning_trades_short,
                        "WINNING PERCENTAGE SHORT": win_percentage_short}

    # *********************************************************************************************************

    def backtesting_all_pair_w_combos(self, num_of_candles, interval,
                                      price_gap, name_of_the_file_to_save):

        import docx

        combos = []

        if interval == "H1":
            for x in range(120, 150 + 1, 5):
                for y in range(120, 150 + 1, 5):
                    combos.append((x, -y))
            tp_sl_gap = 10  # HOW MANY PIPS IN BETWEEN TP/SL BEFORE THE SCRIPT MOVES TO ANOTHER COMBINATION
        elif interval == "H4":
            for x in range(120, 240 + 1, 10):
                for y in range(120, 240 + 1, 10):
                    combos.append((x, -y))
            tp_sl_gap = 20
        elif interval == "M30":
            for x in range(80, 160 + 1, 5):
                for y in range(80, 160 + 1, 5):
                    combos.append((x, -y))
            tp_sl_gap = 10
        elif interval == "M15":
            for x in range(60, 120 + 1, 5):
                for y in range(60, 120 + 1, 5):
                    combos.append((x, -y))
            tp_sl_gap = 10
        else:  # THIS REPRESENTS THE DAILY TIME FRAME
            for x in range(100, 300 + 1, 10):
                for y in range(100, 300 + 1, 10):
                    combos.append((x, -y))
            tp_sl_gap = 30

        self.results_w_combos = []

        for el in combos:

            if el[0] >= -el[1] and el[0] + el[1] <= tp_sl_gap:  # PLUS AS el[1] IS NEGATIVE

                print(el[0], el[1], price_gap)

                self.from_zero_trades_all_pairs = 0
                self.from_zero_trades_w_profit_all_pairs = 0
                self.from_zero_trades_w_loss_all_pairs = 0
                self.from_zero_profit_percentage_all_pairs = 0

                self.from_hundred_trades_all_pairs = 0
                self.from_hundred_trades_w_profit_all_pairs = 0
                self.from_hundred_trades_w_loss_all_pairs = 0
                self.from_hundred_profit_percentage_all_pairs = 0

                self.results = {}

                try:

                    self.backtesting_all_pair(num_of_candles, interval, el[0], el[1], price_gap)

                    self.results_w_combos.append(self.results)

                except:

                    continue
            else:

                continue

        results_w_combos_str = str()

        for el in self.results_w_combos:
            if el:
                results_w_combos_str += str(el)
            else:
                continue

        d = docx.Document()  # ? NOT PASSING ANYTHING WILL CREATE A BLANK DOC
        d.add_paragraph(results_w_combos_str)
        # INSERT YOUR PATH TO THE FOLDER WHERE YOU WANT TO SAVE THE RESULTS
        path = "C:\\Users\\You\\Desktop\\Forex_Bot_Results\\" + name_of_the_file_to_save
        d.save(path)


# *********************************** RUN ONE TIME WITHOUT SAVING THE RESULTS ****************************************

# s = Strategy()

# s.backtesting_single_pair("CAD_SGD", 100, "H1", 125, -125, 5)

# s.backtesting_all_pair(5000, "H1", 140, -140, 5)

# print(s.results, s.in_depth_res_short, s.in_depth_res_long)


# *********************************** RUN MULTIPLE TIME FRAMES / COMBINATIONS ***************************************

def run_them():
    s = Strategy()

    # s.backtesting_all_pair_w_combos(5000, "M15", 2.5, "The_name_of_your_strategy.docx")

    # s.backtesting_all_pair_w_combos(5000, "M15", 5, "The_name_of_your_strategy.docx")

    # s.backtesting_all_pair_w_combos(5000, "H1", 5, "The_name_of_your_strategy.docx")

    # s.backtesting_all_pair_w_combos(5000, "M30", 5, "The_name_of_your_strategy.docx")

    # s.backtesting_all_pair_w_combos(5000, "H4", 5, "The_name_of_your_strategy.docx")

    # s.backtesting_all_pair_w_combos(5000, "H1", 10, "The_name_of_your_strategy.docx")

    # s.backtesting_all_pair_w_combos(833, "H4", 10, "The_name_of_your_strategy.docx")

    # s.backtesting_all_pair_w_combos(3000, "D", 10, "The_name_of_your_strategy.docx")

    # s.backtesting_all_pair_w_combos(5000, "H1", 15, "The_name_of_your_strategy.docx")

    # s.backtesting_all_pair_w_combos(5000, "H4", 15, "The_name_of_your_strategy.docx")

    # s.backtesting_all_pair_w_combos(5000, "H1", 20, "The_name_of_your_strategy.docx")

    # s.backtesting_all_pair_w_combos(5000, "H1", 25, "The_name_of_your_strategy.docx")

    # s.backtesting_all_pair_w_combos(5000, "H4", 25, "The_name_of_your_strategy.docx")

    # s.backtesting_all_pair_w_combos(3000, "D", 25, "The_name_of_your_strategy.docx")

    # s.backtesting_all_pair_w_combos(3000, "D", 50, "The_name_of_your_strategy.docx")

    return "DONE"

# run_them() # REMOVE # IF YOU NEED TO RUN THE FUNCTION


# *********************************** RUN ONLY INTEREST BITS *********************************************

def temp_check():
    import docx

    s = Strategy()

    s.backtesting_all_pair(5000, "H1", 125, -125, 5)

    temp_res_str = str(s.in_depth_res_short)

    d = docx.Document()
    d.add_paragraph(temp_res_str)
    path = 'c:\\users\\wtf1\\desktop\\forex_bot_results\\single_results\\' + \
           "The_name_of_your_strategy.docx"
    d.save(path)

# temp_check() # REMOVE # IF YOU NEED TO RUN THE FUNCTION


# ******************************************** END OF THE SCRIPT ****************************************************
